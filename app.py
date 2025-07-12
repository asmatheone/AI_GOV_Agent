

import streamlit as st
import openai
import os
from PyPDF2 import PdfReader
import docx
from docx import Document



# إعدادات الصفحة
st.set_page_config(page_title="مساعد الحكومي ", page_icon="logo_gov.png", layout="centered")

# CSS لتنسيق الواجهة
st.markdown(
    """
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    .stApp {
        background-color: #ffffff;
        direction: rtl;
    }
    h1, h2, h3 {
        color: #004225 !important;
        text-align: center !important;
    }
    label, .stTextInput, .stTextArea, .stSelectbox {
        text-align: right !important;
        direction: rtl;
    }
    .warning {
        color: black;
        font-weight: bold;
        text-align: center;
        margin-top: -10px;
        margin-bottom: 20px;
    }
    </style>
    """,
    unsafe_allow_html=True
)


# الشعار والعنوان والتحذير
#st.image("logo.png", width=120)

st.markdown("## المساعد الحكومي")

st.markdown('<div class="warning">يرجى الالتزام بسياسات خصوصية البيانات وحماية البيانات الشخصية المعتمدة من الهيئه السعودية للبيانات والذكاء والاصطناعي  </div>', unsafe_allow_html=True)

# إعداد المفتاح
api_key = st.secrets["OPENAI_API_KEY"]
openai.api_key = api_key


# رسائل النظام المختلفة
system_message_general = (
    "أنت مساعد ذكي رسمي يعمل في جهة حكومية . أجب فقط على الأسئلة المتعلقة بالجهات الحكومية وإداراتها، بما في ذلك وليس حصرًا: إدارة الأمن السيبراني، إدارة حوكمة البيانات، الإدارة التقنية، الإدارة العامة للحوكمة والمخاطر، الهيكلة الإدارية، الاستراتيجيات، الإجراءات المالية، وكل ما يتعلق بالمهام الرسمية للوزارات. تجاهل أي أسئلة لا ترتبط مباشرة بالجهات الحكومية ."
)

system_message_email = (
    "أنت مساعد متخصص في كتابة رسائل بريد إلكتروني رسمية نيابة عن جهة حكومية . "
    "قم بصياغة بريد إلكتروني رسمي ومهني ومفصل باللغة العربية بناءً على التعليمات التالية. "
    "ابدأ بتحية رسمية، ثم صغ الرسالة بوضوح واحترافية."
)

system_message_policy = (
    "أنت خبير متخصص في صياغة السياسات الرسمية داخل جهة حكومية . "
    "قم بصياغة سياسة داخلية مفصلة وطويلة وفق الأسلوب الرسمي المعتمد في الحوكمة والالتزام. "
    "تتضمن السياسة: الغرض، النطاق، التعاريف، المسؤوليات، الإجراءات، العقوبات (إن وجدت)، والمراجع التنظيمية. "
    "استخدم لغة رسمية دقيقة واحترافية   ."
)

system_message_risk = (
    "أنت محلل مختص في إدارة المخاطر في جهة حكومية. "
    "مهمتك هي تحليل مؤشرات الأداء وتقديم قائمة دقيقة بالمخاطر المحتملة التي قد تواجهها الجهات الحكومية بناءً على المؤشر المدخل. "
    "لا تجب على أي أسئلة أخرى. قدم فقط المخاطر المتوقعة بشكل واضح، رسمي، ومفصل."
)


# واجهة الخيارات
task = st.selectbox("اختر نوع المهمة", [
    "مساعد ذكي - أجب عن سؤال",
    "تلخيص مستند",
    "مساعدة في كتابة بريد إلكتروني",
    "صياغة سياسة جديدة",
    "المخاطر المتوقعة للمؤشرات"
])

# المهام
if task == "المخاطر المتوقعة للمؤشرات":
    indicator = st.text_area("اكتب اسم المؤشر أو وصفه")
    if st.button("تحليل المخاطر المتوقعة"):
        if indicator:
            with st.spinner("جاري تحليل المؤشر وتحديد المخاطر..."):
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": system_message_risk},
                        {"role": "user", "content": indicator}
                    ]
                )
                risks = response.choices[0].message["content"] 
                st.markdown("### المخاطر المتوقعة:")
                st.write(risks)

elif task == "تلخيص مستند":
    uploaded_file = st.file_uploader("ارفع مستند PDF أو Word", type=["pdf", "docx"])
    if uploaded_file:
        text = ""
        if uploaded_file.name.endswith(".pdf"):
            reader = PdfReader(uploaded_file)
            for page in reader.pages:
                text += page.extract_text()
        elif uploaded_file.name.endswith(".docx"):
            doc = docx.Document(uploaded_file)
            for para in doc.paragraphs:
                text += para.text + "\n"

        if text:
            with st.spinner("جاري التلخيص..."):
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": system_message_general + " قم بتلخيص المستند التالي بشكل شامل ومفصل:"},
                        {"role": "user", "content": text}
                    ]
                )
                summary = response.choices[0].message["content"] 
                st.markdown("### الملخص:")
                st.write(summary)

                docx_file = Document()
                docx_file.add_heading("ملخص المستند", 0)
                docx_file.add_paragraph(summary)
                docx_path = "ملخص_المستند.docx"
                docx_file.save(docx_path)
                with open(docx_path, "rb") as f:
                    st.download_button("تحميل الملخص كـ Word", f, file_name=docx_path)

elif task == "مساعدة في كتابة بريد إلكتروني":
    email_goal = st.text_area("ما الغرض من البريد الإلكتروني؟")
    if st.button("كتابة البريد"):
        if email_goal:
            with st.spinner("جاري إعداد البريد الرسمي..."):
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": system_message_email},
                        {"role": "user", "content": email_goal}
                    ]
                )
                email = response.choices[0].message["content"]
                st.markdown("### البريد الإلكتروني المقترح:")
                st.write(email)

elif task == "صياغة سياسة جديدة":
    policy_topic = st.text_area("ما موضوع السياسة التي ترغب بصياغتها؟")
    if st.button("كتابة السياسة"):
        if policy_topic:
            with st.spinner("جاري صياغة السياسة بالتفصيل..."):
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": system_message_policy},
                        {"role": "user", "content": policy_topic}
                    ]
                )
                policy = response.choices[0].message["content"]
                st.markdown("### السياسة المقترحة:")
                st.write(policy)

                docx_file = Document()
                docx_file.add_heading("السياسة المقترحة", 0)
                docx_file.add_paragraph(policy)
                docx_path = "السياسة_المقترحة.docx"
                docx_file.save(docx_path)
                with open(docx_path, "rb") as f:
                    st.download_button("تحميل السياسة كـ Word", f, file_name=docx_path)

else:
    user_input = st.text_area("اكتب سؤالك هنا")
    if st.button("إرسال"):
        if user_input:
            with st.spinner("جاري إعداد إجابة تفصيلية..."):
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": system_message_general},
                        {"role": "user", "content": user_input}
                    ]
                )
                answer = response.choices[0].message["content"] 
                st.markdown("### الإجابة:")
                st.write(answer)


# ✅ رسالة نهاية الصفحة
# st.markdown('<div class="custom-footer">في حال واجهتك مشكلة يمكنك التواصل على البريد الالكتروني asma.ha.almalki@gmail.com</div>', unsafe_allow_html=True)


# زر مساعدة تفاعلي
with st.expander("📩 هل تحتاج إلى مساعدة؟", expanded=False):
    st.info("في حال وجود استفسار أو واجهتكم مشكلة تقنية، الرجاء التواصل  على البريد الإلكتروني: asma.ha.almalki@gmail.com")
